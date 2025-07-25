import os
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document

# Paths
DATA_PATH = ".\\Doc"  # Directory containing .docx files
CHROMA_PATH = "chroma"  # Relative path for Chroma database

def load_docx_data(docx_path):
    """Load text from a .docx file."""
    try:
        doc = DocxDocument(docx_path)
        # Extract text from paragraphs
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text
    except Exception as e:
        print(f"Error reading {docx_path}: {e}")
        return ""

def collect_docx_files(directory):
    """Collect all .docx files from the specified directory."""
    docx_files = []
    for root, _, files in os.walk(directory):
        for file in files:
            if file.lower().endswith(".docx"):
                docx_files.append(os.path.join(root, file))
    return docx_files

def docx_to_documents(docx_paths):
    """Convert .docx files into LangChain Document objects."""
    documents = []
    
    for docx_path in docx_paths:
        # Load text from .docx file
        content = load_docx_data(docx_path)
        if not content:
            continue
        
        # Create metadata
        file_name = os.path.basename(docx_path)
        file_size = os.path.getsize(docx_path)
        metadata = {
            "file_path": docx_path,
            "file_name": file_name,
            "file_size": file_size
        }
        
        # Create LangChain Document
        doc = Document(page_content=content, metadata=metadata)
        documents.append(doc)
    
    return documents

def split_text(documents):
    """Split documents into chunks using RecursiveCharacterTextSplitter."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=500,
        length_function=len,
        add_start_index=True
    )
    chunks = text_splitter.split_documents(documents)
    return chunks

def create_vec_db(chroma_path, chunks):
    """Create and persist a Chroma vector database from document chunks."""
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    
    db = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=chroma_path
    )
    db.persist()  # Save the database to disk
    return db

def main():
    # Collect all .docx files from the directory
    docx_paths = collect_docx_files(DATA_PATH)
    if not docx_paths:
        print("No .docx files found in the specified directory.")
        return
    
    # Convert .docx files to documents
    documents = docx_to_documents(docx_paths)
    if not documents:
        print("No valid documents created from .docx files.")
        return
    
    # Split documents into chunks
    chunks = split_text(documents)
    
    # Create vector database
    db = create_vec_db(CHROMA_PATH, chunks)
    
    print(f"Vector database created at {CHROMA_PATH} with {len(chunks)} chunks.")

if __name__ == "__main__":
    main()